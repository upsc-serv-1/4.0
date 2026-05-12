"""JSON Tool — FastAPI backend.

PDF → Quiz JSON Extractor (Gemini-assisted, copy-paste flow). Schema 2.0 output.
All routes prefixed `/api`. Uses MongoDB (jt_jobs, jt_questions, jt_revisions, jt_batches).
"""
from __future__ import annotations

import io
import logging
import os
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional, Dict, Any

from dotenv import load_dotenv
from fastapi import FastAPI, APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import Response, JSONResponse, StreamingResponse
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field
from starlette.middleware.cors import CORSMiddleware

from services.pdf_extract import (
    extract_pages, is_scanned, repeated_line_strip, clean_lines, full_text, render_page_png,
)
from services.q_splitter import split_questions, bundle_qp_sol, chunk_into_batches
from services.prompt_builder import build_batch_prompt, load_taxonomy, build_reverify_prompt
from services.output_parser import parse_output, validate_against_taxonomy
from services.exporter import build_schema2_json, build_markdown
from services.docx_export import prompt_to_docx_bytes


ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

UPLOADS_DIR = ROOT_DIR / "data" / "uploads"
UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

mongo_url = os.environ["MONGO_URL"]
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ["DB_NAME"]]

app = FastAPI(title="JSON Tool API")
api = APIRouter(prefix="/api")

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s: %(message)s")
logger = logging.getLogger("jsontool")

# ─────────── INSTITUTE / PROGRAM PATTERNS ────────────────────────────────

INSTITUTE_PATTERNS = [
    ("forum_ias_official", "Forum IAS", ["forum ias", "forumias"]),
    ("pmfias_official", "PMF IAS", ["pmf ias", "pmfias"]),
    ("vision_ias_official", "Vision IAS", ["vision ias", "visionias"]),
    ("drishti_official", "Drishti IAS", ["drishti ias", "drishtiias", "drishti"]),
    ("upsc_official", "UPSC", ["upsc"]),
    ("insights_official", "Insights IAS", ["insights ias", "insightsias"]),
    ("gsbg_official", "GSBG", ["gsbg"]),
]

PROGRAM_PATTERNS = [
    ("gs-simulator", "GS Simulator", ["gs simulator", "gs-simulator"]),
    ("gsbg", "GSBG", ["gsbg"]),
    ("pyq-toolkit", "PYQ Toolkit", ["pyq toolkit", "pyq-toolkit"]),
    ("pyq-book", "PYQ Book", ["pyq book", "pyq-book"]),
    ("test-series", "Test Series", ["test series", "test-series"]),
    ("cse", "CSE", ["cse prelims", "cse mains"]),
]


def auto_fill_from_filename(filename: str) -> Dict[str, Any]:
    """Extract institute, program, and title hints from filename."""
    name = Path(filename).stem
    lower = name.lower()
    result: Dict[str, Any] = {}

    # Institute
    for inst_id, inst_name, patterns in INSTITUTE_PATTERNS:
        if any(p in lower for p in patterns):
            result["institute"] = inst_name
            result["institute_id"] = inst_id
            result["institute_name"] = inst_name
            break

    # Program
    for prog_id, prog_name, patterns in PROGRAM_PATTERNS:
        if any(p in lower for p in patterns):
            result["program_id"] = prog_id
            result["program_name"] = prog_name
            break

    # Title suggestion — clean up the filename
    title = re.sub(r'\.(pdf|PDF)$', '', name)
    title = re.sub(r'[_-]+', ' ', title).strip()
    result["title_suggestion"] = title

    return result


# ───────────────────────────── Models ────────────────────────────────────────

class Metadata(BaseModel):
    id: Optional[str] = None
    title: str
    launch_year: Optional[int] = None
    institute: Optional[str] = ""
    program_id: Optional[str] = ""
    program_name: Optional[str] = ""
    series: Optional[str] = ""
    level: Optional[str] = ""
    paperType: Optional[str] = ""
    defaultMinutes: Optional[int] = None
    sourceMode: Optional[str] = "docx-inline"
    schema_version: Optional[str] = "2.0"
    institute_id: Optional[str] = None
    institute_name: Optional[str] = None
    exam_frame: Dict[str, Any] = Field(
        default_factory=lambda: {
            "exam_category": "cse",
            "specific_exam": None,
            "stage": "prelims",
            "paper": "pre_gs1",
        }
    )


class JobCreateResponse(BaseModel):
    id: str
    title: str
    status: str
    total_questions: int
    created_at: str


class GeneratePromptsRequest(BaseModel):
    batch_size: int = 35
    subject_filter: List[str] = []
    extra_instructions: str = ""
    use_ocr: bool = False
    columns: int = 1


class ParseOutputRequest(BaseModel):
    output_text: str
    batch_index: Optional[int] = None


class QuestionUpdate(BaseModel):
    subject: Optional[str] = None
    section_group: Optional[str] = None
    microtopic: Optional[str] = None
    statement_lines: Optional[List[str]] = None
    options: Optional[Dict[str, str]] = None
    correct_answer: Optional[str] = None
    explanation_markdown: Optional[str] = None
    # PYQ fields
    pyq_source: Optional[str] = None
    pyq_year: Optional[int] = None
    is_pyq: Optional[bool] = None
    is_ncert: Optional[bool] = None
    pyq_exam_label: Optional[str] = None   # "Prelims" / "Mains"
    pyq_group: Optional[str] = None        # "UPSC CSE" / "UPSC CDS" / "BPSC" …
    source_attribution_label: Optional[str] = None
    confidence: Optional[int] = None
    inconsistency_flag: Optional[str] = None
    inconsistency_reason: Optional[str] = None


class BulkQuestionUpdate(BaseModel):
    question_numbers: List[int]
    updates: Dict[str, Any]


class ExportPdfRequest(BaseModel):
    font_family: str = "sans"
    font_size: int = 12
    columns: int = 1
    theme: str = "modern"
    paper_style: str = "plain"
    content_scope: str = "q_options_expl"  # q_only | q_options | q_options_expl
    answer_placement: str = "inline"       # inline | end
    visual_style: str = "document"         # document | flashcard
    qa_background_color: str = "transparent"
    show_toc: bool = False
    header_text: str = ""
    footer_text: str = ""
    watermark: str = ""


# ────────────────────────── Helpers ──────────────────────────────────────────

def slugify(s: str) -> str:
    s = (s or "").lower().strip()
    s = re.sub(r"[^a-z0-9]+", "-", s)
    return re.sub(r"-+", "-", s).strip("-") or "untitled"


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


async def _get_job(job_id: str) -> Dict:
    job = await db.jt_jobs.find_one({"id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(404, f"Job {job_id} not found")
    return job


def _compute_pyq_flags(group: str) -> Dict[str, bool]:
    """Compute is_upsc_cse / is_allied / is_others from group string."""
    g = (group or "").strip().upper()
    is_upsc_cse = g == "UPSC CSE"
    is_allied = g.startswith("UPSC") and g != "UPSC CSE"
    is_others = bool(g) and not g.startswith("UPSC")
    return {"is_upsc_cse": is_upsc_cse, "is_allied": is_allied, "is_others": is_others}


# ─────────────────────────── Routes ──────────────────────────────────────────

@api.get("/")
async def root():
    return {"service": "json-tool", "status": "ok", "version": "0.2.0"}


@api.get("/taxonomy")
async def get_taxonomy():
    return {"entries": load_taxonomy()}


@api.get("/filename-hints")
async def filename_hints(filename: str):
    """Return auto-fill hints from a filename."""
    return auto_fill_from_filename(filename)


@api.get("/jobs")
async def list_jobs():
    cursor = db.jt_jobs.find({}, {"_id": 0}).sort("created_at", -1)
    items = await cursor.to_list(length=500)
    return {"items": items}


@api.post("/jobs", response_model=JobCreateResponse)
async def create_job(
    title: str = Form(...),
    metadata_json: str = Form(...),
    qp_pdf: UploadFile = File(...),
    sol_pdf: Optional[UploadFile] = File(None),
):
    import json as _json
    try:
        metadata = _json.loads(metadata_json)
    except _json.JSONDecodeError as e:
        raise HTTPException(400, f"Invalid metadata_json: {e}")

    if not metadata.get("id"):
        metadata["id"] = slugify(metadata.get("title") or title)

    job_id = str(uuid.uuid4())
    job_dir = UPLOADS_DIR / job_id
    job_dir.mkdir(parents=True, exist_ok=True)

    qp_path = job_dir / "qp.pdf"
    qp_path.write_bytes(await qp_pdf.read())
    sol_path = None
    if sol_pdf is not None:
        # FastAPI may pass empty UploadFile if user doesn't attach; check filename
        if getattr(sol_pdf, "filename", None):
            sp = job_dir / "sol.pdf"
            sp.write_bytes(await sol_pdf.read())
            sol_path = str(sp)

    job_doc = {
        "id": job_id,
        "title": title,
        "status": "created",
        "metadata": metadata,
        "qp_pdf_path": str(qp_path),
        "sol_pdf_path": sol_path,
        "total_questions": 0,
        "batch_size": 35,
        "subject_filter": [],
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    await db.jt_jobs.insert_one(dict(job_doc))
    job_doc.pop("_id", None)
    return JobCreateResponse(
        id=job_id, title=title, status="created", total_questions=0, created_at=job_doc["created_at"]
    )


@api.get("/jobs/{job_id}")
async def get_job(job_id: str):
    job = await _get_job(job_id)
    qs = await db.jt_questions.find({"job_id": job_id}, {"_id": 0}).sort("question_number", 1).to_list(length=2000)
    batches = await db.jt_batches.find({"job_id": job_id}, {"_id": 0}).sort("batch_index", 1).to_list(length=200)
    return {"job": job, "questions": qs, "batches": batches}


@api.delete("/jobs/{job_id}")
async def delete_job(job_id: str):
    await _get_job(job_id)
    await db.jt_questions.delete_many({"job_id": job_id})
    await db.jt_batches.delete_many({"job_id": job_id})
    await db.jt_revisions.delete_many({"job_id": job_id})
    await db.jt_jobs.delete_one({"id": job_id})
    job_dir = UPLOADS_DIR / job_id
    if job_dir.exists():
        for p in job_dir.rglob("*"):
            if p.is_file():
                p.unlink()
        for p in sorted(job_dir.rglob("*"), reverse=True):
            if p.is_dir():
                p.rmdir()
        job_dir.rmdir()
    return {"deleted": True, "id": job_id}


@api.get("/jobs/{job_id}/preview")
async def preview_job(job_id: str, use_ocr: bool = False, columns: int = 1):
    """Extract PDFs, split into questions, return bundle + sanity report."""
    job = await _get_job(job_id)
    try:
        qp_pages = extract_pages(job["qp_pdf_path"], use_ocr=use_ocr, columns=columns)
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
        
    qp_pages = repeated_line_strip(qp_pages)
    qp_text = clean_lines(full_text(qp_pages))
    qp_blocks = split_questions(qp_text)

    sol_blocks: list = []
    sol_pages: list = []
    if job.get("sol_pdf_path"):
        sol_pages = extract_pages(job["sol_pdf_path"], use_ocr=use_ocr, columns=columns)
        sol_pages = repeated_line_strip(sol_pages)
        sol_text = clean_lines(full_text(sol_pages))
        sol_blocks = split_questions(sol_text)

    bundle = bundle_qp_sol(qp_blocks, sol_blocks)
    qp_scanned = is_scanned(qp_pages)
    sol_scanned = is_scanned(sol_pages) if sol_pages else False

    await db.jt_jobs.update_one(
        {"id": job_id},
        {"$set": {
            "total_questions": len(bundle["items"]),
            "status": "extracted",
            "updated_at": now_iso(),
        }},
    )
    return {
        "items": bundle["items"][:200],  # preview cap
        "items_count": len(bundle["items"]),
        "qp_numbers": bundle["qp_numbers"],
        "sol_numbers": bundle["sol_numbers"],
        "missing_in_qp": bundle["missing_in_qp"],
        "missing_in_sol": bundle["missing_in_sol"],
        "total_qp": bundle["total_qp"],
        "total_sol": bundle["total_sol"],
        "qp_scanned": qp_scanned,
        "sol_scanned": sol_scanned,
        "qp_pages": len(qp_pages),
        "sol_pages": len(sol_pages),
    }


@api.post("/jobs/{job_id}/prompts")
async def generate_prompts(job_id: str, body: GeneratePromptsRequest):
    job = await _get_job(job_id)
    try:
        qp_pages = extract_pages(job["qp_pdf_path"], use_ocr=body.use_ocr, columns=body.columns)
    except RuntimeError as e:
        raise HTTPException(status_code=400, detail=str(e))
        
    qp_pages = repeated_line_strip(qp_pages)
    qp_text = clean_lines(full_text(qp_pages))
    qp_blocks = split_questions(qp_text)
    sol_blocks: list = []
    if job.get("sol_pdf_path"):
        sp = extract_pages(job["sol_pdf_path"], use_ocr=body.use_ocr, columns=body.columns)
        sp = repeated_line_strip(sp)
        sol_blocks = split_questions(clean_lines(full_text(sp)))
    bundle = bundle_qp_sol(qp_blocks, sol_blocks)
    items = bundle["items"]
    if not items:
        raise HTTPException(400, "No questions detected in PDF")

    batches = chunk_into_batches(items, batch_size=body.batch_size)
    # Wipe old batches for this job
    await db.jt_batches.delete_many({"job_id": job_id})

    out = []
    for idx, batch in enumerate(batches):
        prompt_text = build_batch_prompt(
            batch,
            batch_index=idx,
            total_batches=len(batches),
            subject_filter=body.subject_filter or None,
            extra_instructions=body.extra_instructions,
        )
        nums = [it["number"] for it in batch]
        doc = {
            "id": str(uuid.uuid4()),
            "job_id": job_id,
            "batch_index": idx,
            "question_numbers": nums,
            "prompt_text": prompt_text,
            "parsed": False,
            "parsed_at": None,
            "created_at": now_iso(),
            "char_count": len(prompt_text),
        }
        await db.jt_batches.insert_one(dict(doc))
        doc.pop("_id", None)
        out.append({k: doc[k] for k in ("id","batch_index","question_numbers","char_count","parsed","created_at")})

    await db.jt_jobs.update_one(
        {"id": job_id},
        {"$set": {
            "status": "prompts_generated",
            "batch_size": body.batch_size,
            "subject_filter": body.subject_filter,
            "total_questions": len(items),
            "updated_at": now_iso(),
        }},
    )
    return {"batches": out, "total": len(items), "batch_count": len(batches)}


@api.get("/jobs/{job_id}/prompts/{batch_index}")
async def get_prompt(job_id: str, batch_index: int):
    await _get_job(job_id)
    b = await db.jt_batches.find_one({"job_id": job_id, "batch_index": batch_index}, {"_id": 0})
    if not b:
        raise HTTPException(404, f"Batch {batch_index} not found")
    return b


@api.get("/jobs/{job_id}/prompts/{batch_index}/docx")
async def get_prompt_docx(job_id: str, batch_index: int):
    await _get_job(job_id)
    b = await db.jt_batches.find_one({"job_id": job_id, "batch_index": batch_index}, {"_id": 0})
    if not b:
        raise HTTPException(404, f"Batch {batch_index} not found")
    data = prompt_to_docx_bytes(b["prompt_text"], title=f"Gemini Prompt Batch {batch_index + 1}")
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="prompt-batch-{batch_index + 1}.docx"'},
    )


@api.post("/jobs/{job_id}/parse-output")
async def parse_output_endpoint(job_id: str, body: ParseOutputRequest):
    await _get_job(job_id)
    parsed = parse_output(body.output_text)
    tax = load_taxonomy()
    parsed["questions"] = validate_against_taxonomy(parsed["questions"], tax)

    saved = 0
    for q in parsed["questions"]:
        n = q["number"]
        existing = await db.jt_questions.find_one({"job_id": job_id, "question_number": n}, {"_id": 0})
        if existing:
            await db.jt_revisions.insert_one({
                "id": str(uuid.uuid4()),
                "job_id": job_id,
                "question_number": n,
                "snapshot": existing,
                "source": "gemini",
                "created_at": now_iso(),
            })
        pyq_source = q.get("pyq_source") or ""
        is_pyq = bool(pyq_source)
        pyq_group = q.get("pyq_group") or ""
        pyq_flags = _compute_pyq_flags(pyq_group)
        doc = {
            "id": existing["id"] if existing else str(uuid.uuid4()),
            "job_id": job_id,
            "question_number": n,
            "subject": q.get("subject"),
            "section_group": q.get("section_group"),
            "microtopic": q.get("microtopic"),
            "microtopic_valid": q.get("microtopic_valid", False),
            "microtopic_known": q.get("microtopic_known", False),
            "statement_lines": q.get("statement_lines", []),
            "question_text": q.get("question_text", ""),
            "options": q.get("options", {}),
            "correct_answer": q.get("correct_answer"),
            "explanation_markdown": q.get("explanation_markdown"),
            # PYQ fields
            "pyq_source": pyq_source,
            "pyq_year": q.get("pyq_year"),
            "is_pyq": is_pyq,
            "is_ncert": q.get("is_ncert", False),
            "pyq_exam_label": q.get("pyq_exam_label", ""),
            "pyq_group": pyq_group,
            "source_attribution_label": q.get("source_attribution_label", ""),
            **pyq_flags,
            "confidence": q.get("confidence", 0),
            "inconsistency_flag": q.get("inconsistency_flag", "none"),
            "inconsistency_reason": q.get("inconsistency_reason", ""),
            "edited": False,
            "parsed_from_gemini": True,
            "created_at": existing["created_at"] if existing else now_iso(),
            "updated_at": now_iso(),
        }
        await db.jt_questions.update_one(
            {"job_id": job_id, "question_number": n},
            {"$set": doc},
            upsert=True,
        )
        saved += 1

    if body.batch_index is not None:
        await db.jt_batches.update_one(
            {"job_id": job_id, "batch_index": body.batch_index},
            {"$set": {"parsed": True, "parsed_at": now_iso()}},
        )

    # Refresh job status
    total_parsed = await db.jt_questions.count_documents({"job_id": job_id})
    job = await _get_job(job_id)
    new_status = "reviewed" if total_parsed >= job.get("total_questions", 0) else "partially_parsed"
    await db.jt_jobs.update_one({"id": job_id}, {"$set": {"status": new_status, "updated_at": now_iso()}})

    return {
        "saved": saved,
        "errors": parsed["errors"],
        "skipped": parsed["skipped"],
        "total_parsed_in_job": total_parsed,
    }


@api.patch("/jobs/{job_id}/questions/{q_num}")
async def update_question(job_id: str, q_num: int, body: QuestionUpdate):
    await _get_job(job_id)
    existing = await db.jt_questions.find_one({"job_id": job_id, "question_number": q_num}, {"_id": 0})
    if not existing:
        raise HTTPException(404, f"Question {q_num} not found")
    await db.jt_revisions.insert_one({
        "id": str(uuid.uuid4()),
        "job_id": job_id,
        "question_number": q_num,
        "snapshot": existing,
        "source": "manual",
        "created_at": now_iso(),
    })
    update = {k: v for k, v in body.dict(exclude_unset=True).items() if v is not None}
    update["edited"] = True
    update["updated_at"] = now_iso()
    if "microtopic" in update or "subject" in update or "section_group" in update:
        tax = load_taxonomy()
        triple = (
            update.get("subject", existing.get("subject")),
            update.get("section_group", existing.get("section_group")),
            update.get("microtopic", existing.get("microtopic")),
        )
        update["microtopic_valid"] = triple in {(t["subject"], t["sectionGroup"], t["microTopic"]) for t in tax}
    # Recompute PYQ flags if group changes
    if "pyq_group" in update:
        flags = _compute_pyq_flags(update["pyq_group"])
        update.update(flags)
    # Auto-build source attribution label
    if ("pyq_group" in update or "pyq_year" in update or "pyq_exam_label" in update) and not update.get("source_attribution_label"):
        group = update.get("pyq_group", existing.get("pyq_group", ""))
        year = update.get("pyq_year", existing.get("pyq_year"))
        label = update.get("pyq_exam_label", existing.get("pyq_exam_label", "Prelims"))
        if group and year:
            update["source_attribution_label"] = f"{group} {label} {year}"

    await db.jt_questions.update_one(
        {"job_id": job_id, "question_number": q_num},
        {"$set": update},
    )
    q = await db.jt_questions.find_one({"job_id": job_id, "question_number": q_num}, {"_id": 0})
    return q


@api.get("/jobs/{job_id}/questions/{q_num}/history")
async def get_revision_history(job_id: str, q_num: int):
    revisions = await db.jt_revisions.find(
        {"job_id": job_id, "question_number": q_num}, {"_id": 0}
    ).sort("created_at", -1).to_list(length=100)
    return {"revisions": revisions}


@api.post("/jobs/{job_id}/questions/{q_num}/restore/{rev_id}")
async def restore_revision(job_id: str, q_num: int, rev_id: str):
    rev = await db.jt_revisions.find_one({"id": rev_id}, {"_id": 0})
    if not rev:
        raise HTTPException(404, "Revision not found")
    
    # Overwrite standard entry with the snapshot block contents
    snapshot = rev["snapshot"]
    snapshot.pop("_id", None) # Sanitization
    snapshot["updated_at"] = now_iso()
    snapshot["edited"] = True
    
    await db.jt_questions.update_one(
        {"job_id": job_id, "question_number": q_num},
        {"$set": snapshot}
    )
    return {"ok": True}


@api.patch("/jobs/{job_id}/bulk-questions")
async def bulk_update_questions(job_id: str, body: BulkQuestionUpdate):
    await _get_job(job_id)
    if not body.question_numbers:
        return {"updated": 0}
    
    patch = {k: v for k, v in body.updates.items() if v is not None}
    patch["updated_at"] = now_iso()
    patch["edited"] = True
    
    r = await db.jt_questions.update_many(
        {"job_id": job_id, "question_number": {"$in": body.question_numbers}},
        {"$set": patch}
    )
    return {"updated": r.modified_count}


@api.get("/jobs/{job_id}/questions")
async def list_questions(job_id: str, confidence_lt: Optional[int] = None):
    await _get_job(job_id)
    query: Dict[str, Any] = {"job_id": job_id}
    if confidence_lt is not None:
        query["confidence"] = {"$lt": confidence_lt}
    qs = await db.jt_questions.find(query, {"_id": 0}).sort("question_number", 1).to_list(length=2000)
    return {"items": qs}


@api.get("/jobs/{job_id}/page-image/{page_num}")
async def page_image(job_id: str, page_num: int, source: str = "qp"):
    job = await _get_job(job_id)
    path = job.get("qp_pdf_path") if source == "qp" else job.get("sol_pdf_path")
    if not path:
        raise HTTPException(404, f"No {source.upper()} PDF for job")
    try:
        png = render_page_png(path, page_num)
    except IndexError as e:
        raise HTTPException(404, str(e))
    return Response(content=png, media_type="image/png")


@api.get("/jobs/{job_id}/page-map")
async def get_page_map(job_id: str, columns: int = 1):
    """Return a dict mapping Question Number -> Page Index."""
    job = await _get_job(job_id)
    from services.pdf_extract import extract_pages
    # No need for OCR in mapping, we only care about simple Text extraction for fast header detection
    qp_pages = extract_pages(job["qp_pdf_path"], use_ocr=False, columns=columns)
    mapping = {}
    from services.q_splitter import Q_HEADING_RE
    for page_idx, pg in enumerate(qp_pages):
        text = pg.get("text", "")
        for m in Q_HEADING_RE.finditer(text):
            try:
                n = int(m.group(1))
                if n not in mapping:
                    mapping[n] = page_idx
            except:
                pass
    return mapping


@api.get("/jobs/{job_id}/export")
async def export_job(job_id: str, format: str = "json"):
    job = await _get_job(job_id)
    qs = await db.jt_questions.find({"job_id": job_id}, {"_id": 0}).sort("question_number", 1).to_list(length=2000)
    if format == "json":
        data = build_schema2_json(job, qs)
        return JSONResponse(data)
    if format == "md":
        text = build_markdown(job, qs)
        return Response(content=text, media_type="text/markdown",
                        headers={"Content-Disposition": f'attachment; filename="{job["metadata"].get("id","job")}.md"'})
    if format == "docx":
        data = _build_output_docx(job, qs)
        fname = f'{job["metadata"].get("id","job")}.docx'
        return Response(
            content=data,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{fname}"'},
        )
    raise HTTPException(400, f"Unsupported format: {format}")


@api.post("/jobs/{job_id}/export/pdf")
async def export_pdf(job_id: str, body: ExportPdfRequest):
    """Generate a styled PDF using reportlab. Returns file bytes."""
    job = await _get_job(job_id)
    qs = await db.jt_questions.find({"job_id": job_id}, {"_id": 0}).sort("question_number", 1).to_list(length=2000)
    pdf_bytes = _build_output_pdf(job, qs, body)
    fname = f'{job["metadata"].get("id","job")}.pdf'
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@api.post("/jobs/{job_id}/reverify-prompt")
async def reverify_prompt(job_id: str, threshold: int = 80):
    await _get_job(job_id)
    qs = await db.jt_questions.find(
        {"job_id": job_id, "confidence": {"$lt": threshold}}, {"_id": 0}
    ).sort("question_number", 1).to_list(length=2000)
    if not qs:
        return {"prompt": "", "count": 0}
    items = []
    prev_micros: Dict[int, str] = {}
    job = await _get_job(job_id)
    qp_pages = extract_pages(job["qp_pdf_path"])
    qp_pages = repeated_line_strip(qp_pages)
    qp_blocks = {b["number"]: b["text"] for b in split_questions(clean_lines(full_text(qp_pages)))}
    sol_blocks_map: Dict[int, str] = {}
    if job.get("sol_pdf_path"):
        sp = repeated_line_strip(extract_pages(job["sol_pdf_path"]))
        sol_blocks_map = {b["number"]: b["text"] for b in split_questions(clean_lines(full_text(sp)))}
    for q in qs:
        n = q["question_number"]
        items.append({
            "number": n,
            "qp_text": qp_blocks.get(n, ""),
            "sol_text": sol_blocks_map.get(n),
        })
        if q.get("microtopic"):
            prev_micros[n] = q["microtopic"]
    prompt = build_reverify_prompt(items, prev_micros)
    return {"prompt": prompt, "count": len(items), "question_numbers": [it["number"] for it in items]}


# ─────────── PDF Generator ───────────────────────────────────────────────────

THEME_COLORS = {
    "modern":    {"bg": (255,255,255), "fg": (15,23,42),   "accent": (16,185,129), "rule": (226,232,240)},
    "classic":   {"bg": (255,255,255), "fg": (17,17,17),   "accent": (29,78,216),  "rule": (229,231,235)},
    "sepia":     {"bg": (247,239,225), "fg": (59,42,24),   "accent": (154,52,18),  "rule": (217,199,163)},
    "historical":{"bg": (247,239,225), "fg": (59,42,24),   "accent": (154,52,18),  "rule": (217,199,163)},
    "dark":      {"bg": (11,15,23),    "fg": (229,231,235), "accent": (96,165,250), "rule": (31,41,55)},
}


def _rgb(tup):
    from reportlab.lib.colors import Color
    return Color(tup[0]/255, tup[1]/255, tup[2]/255)


def _build_output_pdf(job: Dict, questions: List[Dict], opts: ExportPdfRequest) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, PageBreak
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor, Color
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    buf = io.BytesIO()
    md = job.get("metadata") or {}
    colors = THEME_COLORS.get(opts.theme, THEME_COLORS["modern"])
    accent_c = _rgb(colors["accent"])
    fg_c = _rgb(colors["fg"])
    rule_c = _rgb(colors["rule"])

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=14*mm, leftMargin=14*mm, topMargin=18*mm, bottomMargin=18*mm,
        title=md.get("title", ""),
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Heading1"],
        textColor=accent_c, fontSize=opts.font_size + 10, spaceAfter=4, alignment=TA_LEFT)
    meta_style  = ParagraphStyle("Meta", parent=styles["Normal"],
        textColor=accent_c, fontSize=opts.font_size - 2, spaceAfter=8)
    q_style     = ParagraphStyle("QStem", parent=styles["Normal"],
        textColor=fg_c, fontSize=opts.font_size, fontName="Helvetica-Bold", spaceAfter=2, leading=opts.font_size*1.4)
    opt_style   = ParagraphStyle("Opt", parent=styles["Normal"],
        textColor=fg_c, fontSize=opts.font_size - 1, leftIndent=10, spaceAfter=1, leading=opts.font_size*1.3)
    ans_style   = ParagraphStyle("Ans", parent=styles["Normal"],
        textColor=accent_c, fontSize=opts.font_size - 1, fontName="Helvetica-Bold", spaceAfter=2)
    expl_style  = ParagraphStyle("Expl", parent=styles["Normal"],
        textColor=fg_c, fontSize=opts.font_size - 2, spaceAfter=4, leading=opts.font_size*1.35)

    story = []

    # Cover / Test metadata block
    story.append(Paragraph(md.get("title", job.get("title", "Test")), title_style))
    meta_parts = []
    if md.get("institute"):
        meta_parts.append(f"Institute: {md['institute']}")
    if md.get("program_name"):
        meta_parts.append(f"Program: {md['program_name']}")
    if md.get("series"):
        meta_parts.append(f"Series: {md['series']}")
    if md.get("launch_year"):
        meta_parts.append(f"Year: {md['launch_year']}")
    if md.get("level"):
        meta_parts.append(f"Level: {md['level']}")
    total_qs = len(questions)
    meta_parts.append(f"Total Questions: {total_qs}")
    if md.get("defaultMinutes"):
        meta_parts.append(f"Duration: {md['defaultMinutes']} min")
    if meta_parts:
        story.append(Paragraph("  |  ".join(meta_parts), meta_style))
    story.append(HRFlowable(width="100%", color=rule_c, thickness=1, spaceAfter=8))

    answer_key = []
    sorted_qs = sorted(questions, key=lambda x: x.get("question_number", 0))

    if opts.visual_style == "flashcard":
        for q in sorted_qs:
            n = q.get("question_number")
            stem = " ".join(q.get("statement_lines") or [q.get("question_text","")]).strip()
            opts_dict = q.get("options") or {}
            correct = (q.get("correct_answer") or "").upper()
            expl = q.get("explanation_markdown") or ""

            q_cell = [Paragraph(f"<b>Q{n}.</b> {stem}", q_style)]
            if opts.content_scope in ("q_options", "q_options_expl"):
                for k in ("a","b","c","d"):
                    if opts_dict.get(k):
                        q_cell.append(Paragraph(f"<b>{k.upper()})</b> {opts_dict[k]}", opt_style))

            a_cell = []
            if opts.answer_placement == "inline":
                if correct:
                    a_cell.append(Paragraph(f"<b>Answer: {correct}</b>", ans_style))
            if opts.content_scope == "q_options_expl" and expl:
                a_cell.append(Paragraph(expl[:800], expl_style))

            tbl = Table([[q_cell, a_cell]], colWidths=["50%","50%"])
            tbl.setStyle(TableStyle([
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("BOX", (0,0), (-1,-1), 0.5, rule_c),
                ("INNERGRID", (0,0), (-1,-1), 0.5, rule_c),
                ("BACKGROUND", (0,0), (0,-1), _rgb(colors["rule"])),
                ("LEFTPADDING", (0,0), (-1,-1), 6),
                ("RIGHTPADDING", (0,0), (-1,-1), 6),
                ("TOPPADDING", (0,0), (-1,-1), 4),
                ("BOTTOMPADDING", (0,0), (-1,-1), 4),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 4))
    else:
        for i, q in enumerate(sorted_qs):
            n = q.get("question_number")
            stem_lines = q.get("statement_lines") or [q.get("question_text","")]
            stem = "<br/>".join(stem_lines).strip()
            opts_dict = q.get("options") or {}
            correct = (q.get("correct_answer") or "").upper()
            expl = q.get("explanation_markdown") or ""
            is_pyq = q.get("is_pyq", False)
            pyq_label = q.get("source_attribution_label") or q.get("pyq_group","")

            story.append(Paragraph(f"<b>Q{n}.</b> {stem}", q_style))
            if pyq_label and is_pyq:
                story.append(Paragraph(f"[{pyq_label}]", meta_style))

            if opts.content_scope in ("q_options","q_options_expl"):
                for k in ("a","b","c","d"):
                    if opts_dict.get(k):
                        story.append(Paragraph(f"<b>{k.upper()})</b> {opts_dict[k]}", opt_style))

            if opts.answer_placement == "inline" and correct:
                story.append(Paragraph(f"<b>Correct Answer: {correct}</b>", ans_style))
            elif opts.answer_placement == "end" and correct:
                answer_key.append((n, correct))

            if opts.content_scope == "q_options_expl" and expl:
                story.append(Spacer(1, 2))
                story.append(Paragraph("<b>Explanation:</b>", ans_style))
                story.append(Paragraph(expl[:2000], expl_style))

            story.append(HRFlowable(width="100%", color=rule_c, thickness=0.5, spaceAfter=4))
            story.append(Spacer(1, 4))

    if answer_key:
        story.append(PageBreak())
        story.append(Paragraph("Answer Key", title_style))
        story.append(HRFlowable(width="100%", color=rule_c, thickness=1, spaceAfter=8))
        for num, ans in answer_key:
            story.append(Paragraph(f"Q{num}: <b>{ans}</b>", opt_style))
        story.append(Spacer(1, 8))

    doc.build(story)
    return buf.getvalue()


def _build_output_docx(job: Dict, questions: List[Dict]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    md = job.get("metadata") or {}
    buf = io.BytesIO()
    document = Document()

    # Title
    title_p = document.add_heading(md.get("title", job.get("title", "Test")), level=1)

    # Test metadata
    meta_parts = []
    if md.get("institute"):
        meta_parts.append(f"Institute: {md['institute']}")
    if md.get("program_name"):
        meta_parts.append(f"Program: {md['program_name']}")
    if md.get("launch_year"):
        meta_parts.append(f"Year: {md['launch_year']}")
    meta_parts.append(f"Total Questions: {len(questions)}")
    if meta_parts:
        p = document.add_paragraph("  |  ".join(meta_parts))
        p.style = document.styles["Body Text"]

    document.add_paragraph("─" * 60)

    sorted_qs = sorted(questions, key=lambda x: x.get("question_number", 0))
    for q in sorted_qs:
        n = q.get("question_number")
        stem_lines = q.get("statement_lines") or [q.get("question_text","")]
        stem = " ".join(stem_lines).strip()
        opts_dict = q.get("options") or {}
        correct = (q.get("correct_answer") or "").upper()
        expl = q.get("explanation_markdown") or ""
        is_pyq = q.get("is_pyq", False)
        pyq_label = q.get("source_attribution_label") or ""

        q_para = document.add_paragraph()
        q_run = q_para.add_run(f"Q{n}. ")
        q_run.bold = True
        q_para.add_run(stem)

        if pyq_label and is_pyq:
            pq_p = document.add_paragraph(f"[{pyq_label}]")
            pq_p.runs[0].italic = True

        for k in ("a","b","c","d"):
            if opts_dict.get(k):
                opt_p = document.add_paragraph(style="List Bullet")
                opt_run = opt_p.add_run(f"{k.upper()}) ")
                opt_run.bold = True
                opt_p.add_run(opts_dict[k])

        if correct:
            ans_p = document.add_paragraph()
            ans_run = ans_p.add_run(f"Correct Answer: {correct}")
            ans_run.bold = True

        if expl:
            expl_p = document.add_paragraph()
            expl_run = expl_p.add_run("Explanation: ")
            expl_run.bold = True
            expl_p.add_run(expl)

        document.add_paragraph("─" * 40)

    document.save(buf)
    return buf.getvalue()


# ──────────────────────── App wiring ─────────────────────────────────────────

app.include_router(api)
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
