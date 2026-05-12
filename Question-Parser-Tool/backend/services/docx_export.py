"""Generate a .docx file from prompt text (for users who prefer to upload to Gemini)."""
from __future__ import annotations
from io import BytesIO
from docx import Document


def prompt_to_docx_bytes(prompt_text: str, title: str = "Gemini Prompt") -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for para in prompt_text.split("\n\n"):
        # split very long paragraphs, otherwise add as-is
        for line in para.splitlines():
            doc.add_paragraph(line)
        doc.add_paragraph("")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
